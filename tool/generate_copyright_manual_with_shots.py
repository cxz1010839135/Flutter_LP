# -*- coding: utf-8 -*-
"""生成软件著作权《软件说明书》：嵌入「领鹏智能说明书_截图原稿备份.docx」界面截图。

要求：
- 界面以截图原稿备份为准
- 正文避免出现 Cursor / 智能对话辅助编程等扩展能力表述
"""

from __future__ import annotations

import re
import shutil
import zipfile
from pathlib import Path

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Cm, Pt, RGBColor

ROOT = Path(__file__).resolve().parent.parent
BACKUP = ROOT / "readme" / "领鹏智能说明书_截图原稿备份.docx"
SHOTS = ROOT / "docs" / "copyright" / "_shots_from_backup"
BLOCKLY_SHOTS = ROOT / "docs" / "copyright" / "_shots_blockly"
MONITOR_SHOTS = ROOT / "docs" / "copyright" / "_shots_monitor"
OUT = ROOT / "docs" / "copyright" / "软件说明书.docx"
OUT_V205 = ROOT / "docs" / "copyright" / "软件说明书_V2.0.5.docx"
OUT_MD = ROOT / "docs" / "copyright" / "软件说明书.md"
V205_SHOTS = ROOT / "docs" / "copyright" / "_shots_v205"

SOFTWARE = "领鹏智能机器人上位机软件"
VERSION = "V2.0.5"
HEADER = f"{SOFTWARE} {VERSION}"


def set_run_font(run, name="宋体", size=12, bold=False, color=None):
    run.font.name = name
    run._element.rPr.rFonts.set(qn("w:eastAsia"), name)
    run.font.size = Pt(size)
    run.font.bold = bold
    if color is not None:
        run.font.color.rgb = color


def add_page_number(paragraph):
    run = paragraph.add_run("第 ")
    set_run_font(run, size=9)
    fld_begin = OxmlElement("w:fldChar")
    fld_begin.set(qn("w:fldCharType"), "begin")
    run._r.append(fld_begin)
    instr = OxmlElement("w:instrText")
    instr.set(qn("xml:space"), "preserve")
    instr.text = "PAGE"
    run._r.append(instr)
    fld_sep = OxmlElement("w:fldChar")
    fld_sep.set(qn("w:fldCharType"), "separate")
    run._r.append(fld_sep)
    run._r.append(OxmlElement("w:t"))
    fld_end = OxmlElement("w:fldChar")
    fld_end.set(qn("w:fldCharType"), "end")
    run._r.append(fld_end)
    run2 = paragraph.add_run(" 页")
    set_run_font(run2, size=9)


def setup_header_footer(doc: Document):
    section = doc.sections[0]
    section.page_height = Cm(29.7)
    section.page_width = Cm(21.0)
    section.top_margin = Cm(2.2)
    section.bottom_margin = Cm(2.2)
    section.left_margin = Cm(2.5)
    section.right_margin = Cm(2.5)

    hp = section.header.paragraphs[0]
    hp.text = HEADER
    hp.alignment = WD_ALIGN_PARAGRAPH.CENTER
    for run in hp.runs:
        set_run_font(run, size=9)

    fp = section.footer.paragraphs[0]
    fp.alignment = WD_ALIGN_PARAGRAPH.CENTER
    add_page_number(fp)


def add_title(doc, text, level=0):
    p = doc.add_paragraph()
    if level == 0:
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = p.add_run(text)
        set_run_font(run, "黑体", 20, True)
        p.paragraph_format.space_after = Pt(12)
    elif level == 1:
        run = p.add_run(text)
        set_run_font(run, "黑体", 15, True)
        p.paragraph_format.space_before = Pt(14)
        p.paragraph_format.space_after = Pt(6)
    else:
        run = p.add_run(text)
        set_run_font(run, "黑体", 13, True)
        p.paragraph_format.space_before = Pt(8)
        p.paragraph_format.space_after = Pt(4)
    return p


def add_para(doc, text, size=11.5, bold=False, center=False, space_after=6):
    p = doc.add_paragraph()
    if center:
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_after = Pt(space_after)
    p.paragraph_format.line_spacing = 1.35
    run = p.add_run(text)
    set_run_font(run, size=size, bold=bold)
    return p


def add_bullets(doc, items):
    for item in items:
        p = doc.add_paragraph(style="List Bullet")
        p.clear()
        run = p.add_run(item)
        set_run_font(run, size=11)
        p.paragraph_format.space_after = Pt(2)


def add_steps(doc, items):
    for i, item in enumerate(items, 1):
        p = doc.add_paragraph()
        p.paragraph_format.space_after = Pt(2)
        run = p.add_run(f"{i}. {item}")
        set_run_font(run, size=11)


def add_caption(doc, text):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_before = Pt(3)
    p.paragraph_format.space_after = Pt(10)
    run = p.add_run(text)
    set_run_font(run, size=10, color=RGBColor(0x55, 0x55, 0x55))


def add_image(doc, path: Path, caption: str, width_cm=14.8):
    if not path.exists():
        add_para(doc, f"[缺少截图: {path.name}]", bold=True)
        return
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run()
    run.add_picture(str(path), width=Cm(width_cm))
    add_caption(doc, caption)


def add_table(doc, headers, rows):
    table = doc.add_table(rows=1 + len(rows), cols=len(headers))
    table.style = "Table Grid"
    for i, h in enumerate(headers):
        cell = table.rows[0].cells[i]
        cell.text = ""
        run = cell.paragraphs[0].add_run(h)
        set_run_font(run, size=10, bold=True)
    for r_i, row in enumerate(rows):
        for c_i, val in enumerate(row):
            cell = table.rows[r_i + 1].cells[c_i]
            cell.text = ""
            run = cell.paragraphs[0].add_run(str(val))
            set_run_font(run, size=10)
    doc.add_paragraph()


def extract_shots() -> dict[int, Path]:
    if not BACKUP.exists():
        raise FileNotFoundError(BACKUP)
    if SHOTS.exists():
        shutil.rmtree(SHOTS)
    SHOTS.mkdir(parents=True)
    with zipfile.ZipFile(BACKUP) as z:
        for name in z.namelist():
            if not name.startswith("word/media/image"):
                continue
            if not name.lower().endswith((".png", ".jpg", ".jpeg")):
                continue
            dest = SHOTS / Path(name).name
            dest.write_bytes(z.read(name))
    mapping = {}
    for p in SHOTS.glob("image*.png"):
        m = re.match(r"image(\d+)\.png$", p.name, re.I)
        if m:
            mapping[int(m.group(1))] = p
    if len(mapping) < 10:
        raise RuntimeError(f"截图不足，仅提取到 {len(mapping)} 张")
    return apply_v205_overrides(mapping)


def apply_v205_overrides(mapping: dict[int, Path]) -> dict[int, Path]:
    """用 V2.0.5 现场截图覆盖原稿中对应图号（无则保留原稿）。"""
    overrides = {
        2: V205_SHOTS / "home_main.png",
        3: V205_SHOTS / "control_continuous.png",
        4: V205_SHOTS / "control_io.png",
        5: V205_SHOTS / "control_joint.png",
        6: V205_SHOTS / "control_gate.png",
        7: V205_SHOTS / "control_linear.png",
        8: V205_SHOTS / "point_library.png",
        11: V205_SHOTS / "clr_zero.png",
        12: V205_SHOTS / "file_config.png",
        13: V205_SHOTS / "file_manage.png",
        14: V205_SHOTS / "config_home.png",
        15: V205_SHOTS / "driver_params.png",
        17: V205_SHOTS / "address_params.png",
    }
    for idx, path in overrides.items():
        if path.exists():
            mapping[idx] = path
            print(f"V2.0.5 override image{idx} <- {path.name}")
    return mapping


def build_docx(img: dict[int, Path]) -> Path:
    doc = Document()
    setup_header_footer(doc)

    for _ in range(2):
        doc.add_paragraph()
    add_title(doc, SOFTWARE, 0)
    add_para(doc, "软件说明书（用户手册）", size=16, bold=True, center=True, space_after=18)
    add_para(doc, f"软件简称：领鹏智能", center=True, space_after=4)
    add_para(doc, f"版本号：{VERSION}", center=True, space_after=4)
    add_para(doc, "著作权人：____________________（待填写）", center=True, space_after=4)
    add_para(doc, "编写日期：____年____月（待填写）", center=True, space_after=18)
    add_para(
        doc,
        "本说明书依据本软件当前界面截图整理，用于介绍功能、运行环境、安装方法与基本操作，"
        "供最终用户、现场工程师及软件著作权登记审查使用。说明书描述常规连接、操控、点位、"
        "文件管理、驱动参数与可视化编程编辑/保存/下发等业务功能。",
        space_after=12,
    )
    doc.add_page_break()

    add_title(doc, "目录", 1)
    for item in [
        "1. 引言",
        "2. 运行环境",
        "3. 安装与卸载",
        "4. 软件概述",
        "5. 连接页",
        "6. 主页",
        "7. 操控页",
        "8. 点位编辑",
        "9. 界面清零",
        "10. 文件配置与文件管理",
        "11. 驱动器参数与地址参数",
        "12. 可视化编程（常规编辑，含分类与搜索界面）",
        "13. 监控功能（主程序 / 寄存器监视 / 特殊寄存器说明）",
        "14. 文件与目录结构",
        "15. 常见问题",
        "16. 附录",
    ]:
        add_para(doc, item, size=12, space_after=3)
    doc.add_page_break()

    # 1
    add_title(doc, "1. 引言", 1)
    add_title(doc, "1.1 编写目的", 2)
    add_para(
        doc,
        "本说明书用于说明「领鹏智能机器人上位机软件」的功能组成、运行环境、安装部署及日常操作流程，"
        "帮助操作人员与工程人员规范使用本软件。",
    )
    add_title(doc, "1.2 软件背景", 2)
    add_para(
        doc,
        "本软件是 ____________________（待填写，公司名称）面向 LP 系列工业机器人控制器开发的上位机应用程序。"
        "软件通过以太网与机器人控制器建立连接，提供连接管理、状态总览、手动操控、点位示教、可视化程序编辑、"
        "文件同步、驱动参数调试等功能，适用于自动化产线的调试、编程与运维。",
    )
    add_para(
        doc,
        "本软件基于 Google Flutter 跨平台框架开发，支持 Windows 桌面与 Android 平板/手机两种主要运行形态。"
        "本说明书界面截图以 Windows 版当前界面为准。",
    )
    add_title(doc, "1.3 术语说明", 2)
    add_table(
        doc,
        ["术语", "说明"],
        [
            ["上位机", "运行本软件的 PC 或平板，相对机器人控制器而言"],
            ["控制器 / 驱控", "领鹏机器人控制柜，通过 IP 访问"],
            ["可视化编程", "基于积木块的图形化程序编辑（工程文件为 XML，配套 RP4）"],
            ["RP4", "机器人程序文本格式，与 XML 工程配套"],
            ["点位 / 点库", "保存的机器人目标姿态或关节坐标"],
            ["门型 / 直线", "操控页中的轨迹移动方式"],
            ["IO", "数字输入 / 输出信号"],
            ["示教", "把当前机器人位置写入指定点位"],
        ],
    )

    # 2
    add_title(doc, "2. 运行环境", 1)
    add_title(doc, "2.1 硬件环境", 2)
    add_table(
        doc,
        ["项目", "最低要求", "推荐配置"],
        [
            ["处理器", "x64 双核 1.5GHz 或 ARM 四核", "Intel i5 / 工控平板 x64"],
            ["内存", "4 GB", "8 GB 及以上"],
            ["硬盘", "500 MB 可用空间", "2 GB 及以上"],
            ["网络", "百兆以太网，与控制器同一网段", "千兆以太网"],
            ["显示器", "1280×720", "1920×1080 横屏"],
        ],
    )
    add_title(doc, "2.2 软件环境", 2)
    add_para(doc, "Windows 版：", bold=True)
    add_bullets(
        doc,
        [
            "操作系统：Windows 10 / Windows 11（64 位）",
            "运行支撑：Microsoft WebView2 运行时（新电脑建议使用含运行时的 Setup 安装包）",
            "安装包：LPRobot-2.0.5-x64-Setup.exe 或 LPRobot-2.0.5-x64.msi",
        ],
    )
    add_para(doc, "Android 版：", bold=True)
    add_bullets(
        doc,
        [
            "操作系统：Android 5.0（API 21）及以上",
            "屏幕方向：横屏",
            "安装包：LPRobot-2.0.5.apk",
        ],
    )
    add_title(doc, "2.3 网络环境", 2)
    add_bullets(
        doc,
        [
            "上位机与控制器须处于同一局域网或可路由的内网",
            "默认控制器 IP 可在连接页修改并自动记忆",
            "通信协议：HTTP（内网），端口 80",
        ],
    )

    # 3
    add_title(doc, "3. 安装与卸载", 1)
    add_title(doc, "3.1 Windows 安装", 2)
    add_steps(
        doc,
        [
            "若本机已安装旧版，请先卸载。",
            "双击 Setup.exe（推荐）或 MSI 安装包，按向导完成安装。",
            "默认安装目录示例：C:\\Program Files\\Lingpeng\\领鹏智能\\。",
            "安装完成后启动「领鹏智能」，主程序为「领鹏智能.exe」。",
            "安装到 Program Files 时，可写数据目录为：%LOCALAPPDATA%\\Lingpeng\\LPRobot\\。",
            "无需以管理员身份运行程序。",
        ],
    )
    add_title(doc, "3.2 Android 安装", 2)
    add_steps(
        doc,
        [
            "将 APK 拷贝至设备或通过 ADB 安装。",
            "若已安装旧版，建议先卸载再安装。",
            "首次运行按系统提示授予必要的存储权限。",
        ],
    )
    add_title(doc, "3.3 卸载", 2)
    add_para(
        doc,
        "Windows：控制面板 → 程序和功能 →「领鹏智能」→ 卸载。"
        "用户数据目录 %LOCALAPPDATA%\\Lingpeng\\LPRobot\\ 如需彻底清除可手动删除。"
        "Android：系统设置 → 应用管理 → 领鹏智能 → 卸载。",
    )

    # 4
    add_title(doc, "4. 软件概述", 1)
    add_para(
        doc,
        "领鹏智能是连接、编程、操控、监控领鹏机器人的统一人机界面。"
        "典型工作流为：连接控制器 → 主页总览 → 进入操控 / 编程 / 监控 / 维护等模块。",
    )
    add_para(doc, "本软件主要界面截图一览（后文章节展开说明）：", bold=True)
    for idx, cap in [
        (1, "图 4-1 连接页"),
        (2, "图 4-2 主页"),
        (3, "图 4-3 操控页"),
        (8, "图 4-4 点位编辑"),
        (13, "图 4-5 文件管理"),
        (15, "图 4-6 驱动器参数"),
    ]:
        add_image(doc, img[idx], cap, width_cm=13.2)

    # 5 连接
    add_title(doc, "5. 连接页", 1)
    add_para(
        doc,
        "启动软件后进入连接页。输入控制器 IP，点击「连接」建立通信；"
        "也可选择跳过连接，仅进行本地可视化程序编辑。",
    )
    add_image(doc, img[1], "图 5-1 连接页")
    add_title(doc, "5.1 操作步骤", 2)
    add_steps(
        doc,
        [
            "确认电脑网卡与控制器在同一网段。",
            "在「控制器IP」输入框填写控制器地址。",
            "点击「连接」，等待提示连接成功。",
            "连接成功后自动进入主页，并开始状态刷新。",
            "若暂不连机、仅需本地编辑程序，可选择跳过连接。",
        ],
    )
    add_para(doc, "说明：上次使用的 IP 会自动记忆。", size=10)

    # 6 主页
    add_title(doc, "6. 主页", 1)
    add_para(
        doc,
        "主页是连接成功后的总览与导航中枢，可查看实时位姿、IO、报警，"
        "并通过左右侧栏进入操控、编程、监控、维护，以及启动 / 停止 / 调速等运行控制。",
    )
    add_image(doc, img[2], "图 6-1 主页总览")
    add_title(doc, "6.1 界面分区", 2)
    add_bullets(
        doc,
        [
            "顶栏：左侧 Logo「领鹏智能」、笛卡尔坐标（X/Y/Z/W）、关节角（J1～J4）、返回",
            "左侧导航：操控、编程、监控、维护",
            "中央：设备示意图与状态展示",
            "右侧运行侧栏：启动、停止、速度百分比、复位等",
            "底栏：设备名、INPUT/OUTPUT 指示、启动状态、报警、连接状态与 IP",
        ],
    )
    add_title(doc, "6.2 常用操作", 2)
    add_steps(
        doc,
        [
            "观察顶栏坐标与底栏连接状态，确认通信正常。",
            "点击左侧「操控 / 编程 / 监控 / 维护」进入对应模块。",
            "使用右侧「启动 / 停止」控制程序运行，通过圆形数值调节速度百分比。",
            "关注底栏「报警」区域；出现报警时先处理安全问题再继续操作。",
        ],
    )

    # 7 操控
    add_title(doc, "7. 操控页", 1)
    add_para(
        doc,
        "操控页用于手动点动、轨迹移动与 IO 操作。左侧选择轴或 IO，中央为速度与模式设置，"
        "右侧切换关节 / 门型 / 直线 / 点位编辑 / 界面清零。",
    )
    add_image(doc, img[3], "图 7-1 操控页（连续模式）")
    add_title(doc, "7.1 点动与速度", 2)
    add_bullets(
        doc,
        [
            "最大速度、加速度：显示当前点动相关参数",
            "速度设定：通过「− / +」或分段条调节速度百分比",
            "模式选择：「连续」为按住持续运动；「长/中/短距离」为固定步进",
            "左侧 X / Y / Z：选择当前操作轴",
        ],
    )
    add_title(doc, "7.2 IO 监视", 2)
    add_para(doc, "在操控页左侧点击「I/O」，可查看多路输入 / 输出通断状态（绿色表示有效）。")
    add_image(doc, img[4], "图 7-2 操控页 · IO 监视")
    add_title(doc, "7.3 关节模式", 2)
    add_para(doc, "右侧点击「关节」，进入关节空间相关点动参数界面。")
    add_image(doc, img[5], "图 7-3 操控页 · 关节模式")
    add_title(doc, "7.4 门型运动", 2)
    add_para(
        doc,
        "右侧点击「门型」，设置目标点、避障高度与速度，确认后按门型轨迹运动到目标点。",
    )
    add_image(doc, img[6], "图 7-4 门型运动参数对话框")
    add_title(doc, "7.5 直线运动", 2)
    add_para(doc, "右侧点击「直线」，设置目标点与速度，确认后直线到达目标点。")
    add_image(doc, img[7], "图 7-5 直线运动参数对话框")

    # 8 点位
    add_title(doc, "8. 点位编辑", 1)
    add_para(
        doc,
        "点位编辑用于管理机器人目标点。列表显示点编号、名称及各轴数值；"
        "右侧提供新增、修改、刷新、删除。",
    )
    add_image(doc, img[8], "图 8-1 点位列表")
    add_title(doc, "8.1 新增点位", 2)
    add_para(doc, "点击「新增」，填写点编号（必填）、名称与描述（可选），确定后写入点库。")
    add_image(doc, img[9], "图 8-2 添加点位对话框")
    add_title(doc, "8.2 当前位置示教", 2)
    add_para(
        doc,
        "可将机器人当前姿态示教到已有点位。系统会弹出确认提示，确认后覆盖该点坐标。"
        "示教前请确认目标点选择正确，避免误覆盖。",
    )
    add_image(doc, img[10], "图 8-3 示教确认提示")

    # 9 清零
    add_title(doc, "9. 界面清零", 1)
    add_para(
        doc,
        "「界面清零」用于按规范进行轴清零。左侧提供清零规范示意与文字说明，"
        "右侧可对各轴执行清零。清零前请仔细阅读机械规范，错误清零可能导致坐标系异常。",
    )
    add_image(doc, img[11], "图 9-1 界面清零")

    # 10 文件
    add_title(doc, "10. 文件配置与文件管理", 1)
    add_title(doc, "10.1 配置首页", 2)
    add_para(
        doc,
        "维护入口进入配置后，可执行机代码自动运行开关、调试模式开关，"
        "并进入文件管理或驱动器参数设置。",
    )
    add_image(doc, img[14], "图 10-1 配置首页")
    add_title(doc, "10.2 文件配置", 2)
    add_para(
        doc,
        "文件配置页左侧提供参数填写说明，中央显示当前配置文件内容。"
        "若提示文件不存在，可先创建该文件后继续编辑，保存后按提示重启驱控。",
    )
    add_image(doc, img[12], "图 10-2 文件配置")
    add_title(doc, "10.3 文件管理", 2)
    add_para(
        doc,
        "文件管理采用左右双栏：左侧为本地目录，右侧为驱控目录。"
        "支持上传到驱控、下载到本地，以及一键备份 / 一键恢复。"
        "上传前请先在右侧进入目标子目录。",
    )
    add_image(doc, img[13], "图 10-3 文件管理（本地 ↔ 驱控）")
    add_steps(
        doc,
        [
            "在左侧选中要上传的本地文件（如 main.xml / main.rp4）。",
            "在右侧进入驱控目标文件夹（不要停在根目录）。",
            "点击「上传到驱控」；或从驱控选中文件后「下载到本地」。",
            "需要整机配置保全时，可使用「一键备份 / 一键恢复」。",
        ],
    )

    # 11 驱动
    add_title(doc, "11. 驱动器参数与地址参数", 1)
    add_title(doc, "11.1 驱动器参数", 2)
    add_para(
        doc,
        "驱动器参数页面向技术人员，用于查看指令/反馈量、设置电机参数、增益与安全阈值，"
        "并对当前轴执行读取 / 写入 / 软复位 / 寻相 / 点动等操作。"
        "请在充分理解参数含义后再写入，避免损坏设备。",
    )
    add_image(doc, img[15], "图 11-1 驱动器参数（运动参数）")
    add_para(
        doc,
        "页面下方可在「运动参数」与「单轴参数」之间切换。"
        "图 11-1 为「运动参数」视图；切换至「单轴参数」可编辑单轴相关参数文件。",
        size=10,
    )
    add_title(doc, "11.2 地址参数", 2)
    add_para(
        doc,
        "地址参数页提供参数控制、总线参数与 SDO 参数的按地址读写，便于底层调试。"
        "写入前请核对地址与数值。可从驱动器参数页点击「地址参数」进入。",
    )
    add_image(doc, img[17], "图 11-2 地址参数")

    # 12 可视化编程（常规编辑，含界面截图）
    # 仅嵌入已核对的 V2.0.5 截图（逻辑总览、搜索）；变量/运动/自定义以文字说明，避免旧版图文错位。
    b = {
        "logic": BLOCKLY_SHOTS / "blockly_logic.png",
        "search": BLOCKLY_SHOTS / "blockly_search.png",
    }
    add_title(doc, "12. 可视化编程（常规编辑）", 1)
    add_para(
        doc,
        "主页点击「编程」进入「领鹏智能编程」页。本软件通过内嵌页面加载本地积木编辑器，"
        "现场工程师可用拖拽方式组合逻辑，生成机器人工程文件（XML / RP4），无需手写底层代码。"
        "左侧工具箱按「逻辑 / 变量 / 运动 / 自定义」分类；中央为工作区；顶栏右侧提供 AI 入口、刷新与返回；"
        "工作区右上角提供搜索、保存、函数库、撤销等工具，右下角提供文件、居中与缩放。",
    )
    add_para(
        doc,
        "本说明书登记范围描述常规可视化编程的编辑、分类积木、搜索定位、保存、导入与下发功能。",
        size=10,
    )
    add_title(doc, "12.1 主要功能", 2)
    add_bullets(
        doc,
        [
            "分类工具箱：逻辑、变量、运动、自定义",
            "拖拽积木组合条件、运算、寄存器与运动指令",
            "加载 / 保存 XML 工程",
            "保存至控制器程序目录（config/server/）",
            "保存至用户工程（files/projects/{名称}/）",
            "保存函数库（files/funlib/）",
            "从本地 XML 目录打开已有工程",
            "积木搜索与结果列表定位（高亮命中块）",
            "退出编程时可按提示将程序上传至控制器在线运行",
        ],
    )
    add_title(doc, "12.2 界面总览与分类积木", 2)
    add_para(
        doc,
        "进入编程页后，左侧可见「逻辑 / 变量 / 运动 / 自定义」四类入口，中央为积木工作区。"
        "「逻辑」分类提供条件判断、与/或运算及比较积木，用于搭建程序分支；"
        "「变量」分类提供寄存器读写与运算积木（X / Y / M / S / T / C / D 等）；"
        "「运动」分类提供自由门型、独立定位、电子齿轮、批量赋值等；"
        "「自定义」分类用于函数块等用户自定义逻辑入口，便于复用子流程。",
    )
    add_image(doc, b["logic"], "图 12-1 可视化编程 · 逻辑分类与工作区")
    add_title(doc, "12.3 搜索与定位", 2)
    add_para(
        doc,
        "在大型工程中，可通过工作区搜索栏输入寄存器号或关键字（如 D8590），"
        "系统列出命中位置，点击后工作区定位并高亮对应积木，便于查阅与修改。",
    )
    add_image(doc, b["search"], "图 12-2 可视化编程 · 搜索定位（示例：D8590）")
    add_title(doc, "12.4 基本操作", 2)
    add_steps(
        doc,
        [
            "在主页点击「编程」，等待编辑器加载完成。",
            "在左侧选择「逻辑 / 变量 / 运动 / 自定义」，从工具箱拖拽积木到工作区并连接。",
            "需要查找寄存器或关键字时，使用搜索栏输入并点击结果定位。",
            "点击保存图标，输入工程名后写入本地程序目录。",
            "需要时可通过导入功能选择已有 XML 追加或加载。",
            "退出编程页时，按提示确认是否上传控制器。",
        ],
    )

    # 13 监控（含界面截图）
    m = {
        "main": MONITOR_SHOTS / "monitor_main.png",
        "special": MONITOR_SHOTS / "monitor_special_regs.png",
    }
    add_title(doc, "13. 监控功能", 1)
    add_para(
        doc,
        "主页左侧进入「监控」后，可监视主程序运行状态与寄存器。"
        "界面上方提供自动运行 / 停止与速度百分比调节；左侧显示主程序（main.rp4）文本；"
        "右侧为寄存器监视；底部为状态窗口，显示运行提示（如 D9000 状态说明）。",
    )
    add_image(doc, m["main"], "图 13-1 监控页（主程序 / 寄存器监视 / 状态窗口）")
    add_title(doc, "13.1 寄存器监视", 2)
    add_bullets(
        doc,
        [
            "按类型切换监视：D / M / S / X / Y",
            "支持添加监视项、删除监视项、自动刷新及刷新周期设置",
            "监视数量有上限（界面显示当前数量/上限），配置可持久化保存",
            "底部状态窗口显示寄存器相关运行提示",
        ],
    )
    add_title(doc, "13.2 特殊寄存器说明", 2)
    add_para(
        doc,
        "在寄存器监视区可打开「特殊寄存器说明」。"
        "D8000～D9999 为设备状态信息或其他特殊用途变量，请勿用作普通地址。"
        "说明对话框列出常用特殊寄存器含义，便于现场查阅。",
    )
    add_bullets(
        doc,
        [
            "D8000：机械手暂停（非 0 减速急停；0 继续运动）",
            "D8001：动态抓取指令阻塞执行策略",
            "D8002：机械手停止（非 0 减速急停；置 0 后当前运动不恢复）",
            "D8004：PLC 状态（0 停止，5 运行）",
            "D8008：非标定制模式类型",
            "D8100～D8199：各轴扭矩/模式相关（只读）",
        ],
    )
    add_image(doc, m["special"], "图 13-2 特殊寄存器说明对话框")

    # 14 目录
    add_title(doc, "14. 文件与目录结构", 1)
    add_para(doc, "安装目录（示例）：", bold=True)
    add_bullets(
        doc,
        [
            "领鹏智能.exe：主程序",
            "data/：运行资源",
            "config/：配置与默认程序（含 config/server/）",
            "dll/visualprogram.lpk：可视化编程资源包",
        ],
    )
    add_para(doc, "用户可写数据目录（Program Files 安装时）：", bold=True)
    add_bullets(
        doc,
        [
            "%LOCALAPPDATA%\\Lingpeng\\LPRobot\\config\\",
            "%LOCALAPPDATA%\\Lingpeng\\LPRobot\\files\\projects\\",
            "%LOCALAPPDATA%\\Lingpeng\\LPRobot\\files\\xml\\",
            "%LOCALAPPDATA%\\Lingpeng\\LPRobot\\files\\funlib\\",
            "%LOCALAPPDATA%\\Lingpeng\\LPRobot\\cache\\visualprogram\\",
        ],
    )

    # 15 FAQ
    add_title(doc, "15. 常见问题", 1)
    add_para(doc, "Q：连接失败怎么办？", bold=True)
    add_para(doc, "A：检查网线/网段、控制器电源与 IP；确认防火墙未拦截本软件访问控制器 80 端口。")
    add_para(doc, "Q：编程页打不开？", bold=True)
    add_para(
        doc,
        "A：Windows 需具备 WebView2 运行时；建议使用 Setup 安装包。"
        "若安装在 Program Files，请确认用户缓存目录可写。",
    )
    add_para(doc, "Q：上传到驱控按钮灰显？", bold=True)
    add_para(doc, "A：请先在右侧驱控目录进入目标子文件夹，并在左侧选中待上传文件。")

    # 16
    add_title(doc, "16. 附录", 1)
    add_para(doc, f"软件名称：{SOFTWARE}")
    add_para(doc, f"版本号：{VERSION}")
    add_para(doc, "著作权人：____________________（待填写）")
    add_para(doc, "开发完成日期：____年____月____日（待填写）")
    add_para(doc, "— 说明书结束 —", center=True, space_after=12)

    OUT.parent.mkdir(parents=True, exist_ok=True)
    written = None
    for path in (OUT_V205, OUT):
        try:
            doc.save(path)
            print("Wrote", path)
            written = path
        except PermissionError:
            print("Skip (locked):", path)
    if written is None:
        alt = OUT.with_name("软件说明书_V2.0.5_另存.docx")
        doc.save(alt)
        print("Wrote", alt, "(原文件被占用，已另存)")
        return alt
    return written


def write_md_summary() -> None:
    text = f"""# {SOFTWARE} {VERSION}

## 软件说明书（用户手册）

**软件名称：** {SOFTWARE}  
**软件简称：** 领鹏智能  
**版本号：** {VERSION}  
**著作权人：** ____________________（待填写）  
**编写日期：** ____年____月（待填写）  

---

本说明书 Word 版由 `tool/generate_copyright_manual_with_shots.py` 生成。

界面截图来源：
- `readme/领鹏智能说明书_截图原稿备份.docx`（连接/主页等基础界面）
- `docs/copyright/_shots_v205/`（V2.0.5 主页/操控/点位/清零/文件/配置/驱动/地址等）
- `docs/copyright/_shots_blockly/`（可视化编程：逻辑总览、搜索定位）
- `docs/copyright/_shots_monitor/`（监控页、特殊寄存器说明）

### 登记范围说明

本说明书描述常规连接、主页、操控、点位、清零、文件配置/管理、驱动参数、地址参数、
可视化编程（分类积木、搜索定位、编辑/保存/导入/下发）及监控（寄存器监视/特殊寄存器说明）等业务功能。

### 生成方法

```powershell
python tool/generate_copyright_manual_with_shots.py
```

生成文件：`docs/copyright/软件说明书.docx`、`docs/copyright/软件说明书_V2.0.5.docx`
"""
    OUT_MD.write_text(text, encoding="utf-8")
    print("Wrote", OUT_MD)


def ensure_blockly_shots() -> None:
    required = [
        "blockly_logic.png",
        "blockly_search.png",
    ]
    missing = [n for n in required if not (BLOCKLY_SHOTS / n).exists()]
    if missing:
        raise FileNotFoundError(
            f"缺少可视化编程截图（请放入 {BLOCKLY_SHOTS}）: {', '.join(missing)}"
        )


def ensure_monitor_shots() -> None:
    required = ["monitor_main.png", "monitor_special_regs.png"]
    missing = [n for n in required if not (MONITOR_SHOTS / n).exists()]
    if missing:
        raise FileNotFoundError(
            f"缺少监控截图（请放入 {MONITOR_SHOTS}）: {', '.join(missing)}"
        )


def main():
    img = extract_shots()
    ensure_blockly_shots()
    ensure_monitor_shots()
    print("shots:", sorted(img))
    print("blockly:", sorted(p.name for p in BLOCKLY_SHOTS.glob("blockly_*.png")))
    print("monitor:", sorted(p.name for p in MONITOR_SHOTS.glob("monitor_*.png")))
    out_path = build_docx(img)
    write_md_summary()
    with zipfile.ZipFile(out_path) as z:
        xml = z.read("word/document.xml").decode("utf-8", errors="ignore")
    for bad in ("Cursor", "cursor.com", "Cursor Agent", "智能对话辅助"):
        if bad in xml:
            raise SystemExit(f"说明书含禁用词: {bad}")
    blip = xml.count("a:blip")
    print(f"OK: no forbidden tool names; embedded images ≈ {blip}")


if __name__ == "__main__":
    main()

#!/usr/bin/env python3
# -*- coding: utf-8 -*-
"""将软件著作权材料 Markdown/TXT 转为 Word（.docx）。"""

from __future__ import annotations

import re
import sys
from pathlib import Path

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_LINE_SPACING
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Cm, Pt, RGBColor

ROOT = Path(__file__).resolve().parent.parent
COPYRIGHT_DIR = ROOT / "docs" / "copyright"
SOFTWARE_FULL = "领鹏智能机器人上位机软件"
VERSION = "V1.7.9"
HEADER_TEXT = f"{SOFTWARE_FULL} {VERSION}"

# 待填写占位符（后续由用户告知后替换或直接在 Word 中填写）
PLACEHOLDER = "____________________（待填写）"
PLACEHOLDER_DATE = "____年____月____日（待填写）"
PLACEHOLDER_MONTH = "____年____月（待填写）"


def set_run_font(run, name: str = "宋体", size: float = 12, bold: bool = False, color=None):
    run.font.name = name
    run._element.rPr.rFonts.set(qn("w:eastAsia"), name)
    run.font.size = Pt(size)
    run.font.bold = bold
    if color:
        run.font.color.rgb = color


def add_header_footer(doc: Document, title: str = HEADER_TEXT):
    section = doc.sections[0]
    section.page_height = Cm(29.7)
    section.page_width = Cm(21.0)
    section.top_margin = Cm(2.54)
    section.bottom_margin = Cm(2.54)
    section.left_margin = Cm(3.17)
    section.right_margin = Cm(3.17)

    header = section.header
    hp = header.paragraphs[0] if header.paragraphs else header.add_paragraph()
    hp.text = title
    hp.alignment = WD_ALIGN_PARAGRAPH.CENTER
    for run in hp.runs:
        set_run_font(run, "宋体", 9)

    footer = section.footer
    fp = footer.paragraphs[0] if footer.paragraphs else footer.add_paragraph()
    fp.alignment = WD_ALIGN_PARAGRAPH.CENTER
    _add_page_number(fp)


def _add_page_number(paragraph):
    run = paragraph.add_run("第 ")
    set_run_font(run, "宋体", 9)
    fld_begin = OxmlElement("w:fldChar")
    fld_begin.set(qn("w:fldCharType"), "begin")
    run._r.append(fld_begin)
    instr = OxmlElement("w:instrText")
    instr.set(qn("xml:space"), "preserve")
    instr.text = "PAGE"
    run._r.append(instr)
    fld_sep = OxmlElement("w:fldChar")
    fld_sep.set(qn("w:fldCharType"), "separate")
    run._r.append(fld_sep)
    run._r.append(OxmlElement("w:t"))
    fld_end = OxmlElement("w:fldChar")
    fld_end.set(qn("w:fldCharType"), "end")
    run._r.append(fld_end)
    run2 = paragraph.add_run(" 页")
    set_run_font(run2, "宋体", 9)


def add_title(doc: Document, text: str, level: int = 0):
    if level == 0:
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = p.add_run(text)
        set_run_font(run, "黑体", 22, bold=True)
        p.paragraph_format.space_after = Pt(12)
    elif level == 1:
        p = doc.add_paragraph()
        run = p.add_run(text)
        set_run_font(run, "黑体", 16, bold=True)
        p.paragraph_format.space_before = Pt(12)
        p.paragraph_format.space_after = Pt(6)
    elif level == 2:
        p = doc.add_paragraph()
        run = p.add_run(text)
        set_run_font(run, "黑体", 14, bold=True)
        p.paragraph_format.space_before = Pt(8)
        p.paragraph_format.space_after = Pt(4)
    else:
        p = doc.add_paragraph()
        run = p.add_run(text)
        set_run_font(run, "黑体", 12, bold=True)
        p.paragraph_format.space_before = Pt(6)
        p.paragraph_format.space_after = Pt(3)


def add_body(doc: Document, text: str, indent: bool = False):
    text = _replace_placeholders(text)
    p = doc.add_paragraph()
    if indent:
        p.paragraph_format.first_line_indent = Cm(0.74)
    p.paragraph_format.line_spacing_rule = WD_LINE_SPACING.MULTIPLE
    p.paragraph_format.line_spacing = 1.5
    _add_mixed_text(p, text)
    return p


def add_bullet(doc: Document, text: str):
    text = _replace_placeholders(text)
    p = doc.add_paragraph(style="List Bullet")
    p.paragraph_format.line_spacing_rule = WD_LINE_SPACING.MULTIPLE
    p.paragraph_format.line_spacing = 1.5
    _add_mixed_text(p, text)


def add_code_block(doc: Document, lines: list[str]):
    for line in lines:
        p = doc.add_paragraph()
        p.paragraph_format.line_spacing = Pt(14)
        p.paragraph_format.space_before = Pt(0)
        p.paragraph_format.space_after = Pt(0)
        run = p.add_run(line if line else " ")
        set_run_font(run, "Courier New", 9)


def _replace_placeholders(text: str) -> str:
    replacements = {
        "【请填写单位或个人名称】": PLACEHOLDER,
        "【请填写】": PLACEHOLDER,
        "2026年06月25日 **【需填写：以实际开发完成日为准】**": PLACEHOLDER_DATE,
        "2026年06月25日": PLACEHOLDER_DATE,
        "2026 年 06 月": PLACEHOLDER_MONTH,
        "广东领鹏智能科技有限公司（领鹏智能）": f"{PLACEHOLDER}（领鹏智能）",
        "广东领鹏智能科技有限公司": PLACEHOLDER,
        "__________________________": PLACEHOLDER,
        "未发表 / 2026年__月__日 **【需填写】**": f"未发表 / {PLACEHOLDER_DATE}",
    }
    for old, new in replacements.items():
        text = text.replace(old, new)
    return text


def _add_mixed_text(paragraph, text: str):
    """支持 **粗体** 与 `代码` 混排。"""
    parts = re.split(r"(\*\*[^*]+\*\*|`[^`]+`)", text)
    for part in parts:
        if not part:
            continue
        if part.startswith("**") and part.endswith("**"):
            run = paragraph.add_run(part[2:-2])
            set_run_font(run, "宋体", 12, bold=True)
        elif part.startswith("`") and part.endswith("`"):
            run = paragraph.add_run(part[1:-1])
            set_run_font(run, "Courier New", 10)
        else:
            run = paragraph.add_run(part)
            set_run_font(run, "宋体", 12)


def add_table(doc: Document, headers: list[str], rows: list[list[str]]):
    table = doc.add_table(rows=1 + len(rows), cols=len(headers))
    table.style = "Table Grid"
    hdr_cells = table.rows[0].cells
    for i, h in enumerate(headers):
        hdr_cells[i].text = ""
        p = hdr_cells[i].paragraphs[0]
        run = p.add_run(_replace_placeholders(h.strip("*")))
        set_run_font(run, "宋体", 10, bold=True)
    for ri, row in enumerate(rows):
        for ci, cell in enumerate(row):
            table.rows[ri + 1].cells[ci].text = ""
            p = table.rows[ri + 1].cells[ci].paragraphs[0]
            _add_mixed_text(p, _replace_placeholders(cell))
    doc.add_paragraph()


def parse_markdown_table(lines: list[str], start: int) -> tuple[list[str], list[list[str]], int]:
    headers = [c.strip() for c in lines[start].strip("|").split("|")]
    rows = []
    i = start + 2
    while i < len(lines) and lines[i].strip().startswith("|"):
        rows.append([c.strip() for c in lines[i].strip("|").split("|")])
        i += 1
    return headers, rows, i


def build_manual_doc() -> Document:
    md_path = COPYRIGHT_DIR / "软件说明书.md"
    lines = md_path.read_text(encoding="utf-8").splitlines()

    doc = Document()
    add_header_footer(doc)

    i = 0
    code_buf: list[str] | None = None

    while i < len(lines):
        line = lines[i]
        stripped = line.strip()

        if code_buf is not None:
            if stripped.startswith("```"):
                add_code_block(doc, code_buf)
                code_buf = None
            else:
                code_buf.append(line)
            i += 1
            continue

        if stripped.startswith("```"):
            code_buf = []
            i += 1
            continue

        if stripped == "---":
            i += 1
            continue

        if stripped.startswith("# ") and "软件说明书" not in stripped:
            add_title(doc, stripped[2:].strip(), 0)
            i += 1
            continue

        if stripped == "## 软件说明书（用户手册）":
            add_title(doc, stripped[3:].strip(), 0)
            i += 1
            continue

        if stripped.startswith("## "):
            add_title(doc, stripped[3:].strip(), 1)
            i += 1
            continue

        if stripped.startswith("### "):
            add_title(doc, stripped[4:].strip(), 2)
            i += 1
            continue

        if stripped.startswith("|") and i + 1 < len(lines) and "---" in lines[i + 1]:
            headers, rows, ni = parse_markdown_table(lines, i)
            add_table(doc, headers, rows)
            i = ni
            continue

        if stripped.startswith("- "):
            add_bullet(doc, stripped[2:])
            i += 1
            continue

        if re.match(r"^\d+\.\s", stripped):
            text = re.sub(r"^\d+\.\s", "", stripped)
            p = doc.add_paragraph()
            p.paragraph_format.line_spacing_rule = WD_LINE_SPACING.MULTIPLE
            p.paragraph_format.line_spacing = 1.5
            _add_mixed_text(p, f"{stripped.split('.')[0]}. { _replace_placeholders(text)}")
            i += 1
            continue

        if stripped.startswith("**") and "：" in stripped:
            add_body(doc, stripped)
            i += 1
            continue

        if stripped and not stripped.startswith("[") and stripped != "**— 说明书结束 —**":
            add_body(doc, stripped)
        elif stripped == "**— 说明书结束 —**":
            p = doc.add_paragraph()
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            run = p.add_run("— 说明书结束 —")
            set_run_font(run, "宋体", 12, bold=True)

        i += 1

    return doc


def build_registration_doc() -> Document:
    md_path = COPYRIGHT_DIR / "登记信息表.md"
    lines = md_path.read_text(encoding="utf-8").splitlines()

    doc = Document()
    add_header_footer(doc, f"{HEADER_TEXT} — 登记信息表")

    add_title(doc, "计算机软件著作权登记信息表", 0)
    add_body(doc, "以下内容供在中国版权保护中心（CPCC）在线登记时填写。标注「待填写」的项请后续补充。")

    i = 0
    while i < len(lines):
        line = lines[i]
        stripped = line.strip()

        if stripped.startswith("## "):
            add_title(doc, stripped[3:].strip(), 1)
            i += 1
            continue

        if stripped.startswith("|") and i + 1 < len(lines) and "---" in lines[i + 1]:
            headers, rows, ni = parse_markdown_table(lines, i)
            add_table(doc, headers, rows)
            i = ni
            continue

        if stripped.startswith("- ") or stripped.startswith(">"):
            text = stripped.lstrip("> ").lstrip("- ")
            add_body(doc, text)
            i += 1
            continue

        if re.match(r"^\d+\.\s", stripped):
            add_body(doc, stripped)
            i += 1
            continue

        if stripped and not stripped.startswith("#") and stripped != "---":
            if not stripped.startswith("[") and "http" not in stripped:
                add_body(doc, stripped)

        i += 1

    return doc


def build_source_doc(txt_name: str, out_name: str) -> Document:
    txt_path = COPYRIGHT_DIR / txt_name
    content = txt_path.read_text(encoding="utf-8").splitlines()

    doc = Document()
    add_header_footer(doc, f"{HEADER_TEXT} — 源程序")

    section = doc.sections[0]
    section.left_margin = Cm(2.5)
    section.right_margin = Cm(2.5)

    for line in content:
        p = doc.add_paragraph()
        p.paragraph_format.line_spacing = Pt(12)
        p.paragraph_format.space_before = Pt(0)
        p.paragraph_format.space_after = Pt(0)
        run = p.add_run(line if line else " ")
        set_run_font(run, "Courier New", 8)

    out_path = COPYRIGHT_DIR / out_name
    doc.save(out_path)
    return out_path


def main():
    COPYRIGHT_DIR.mkdir(parents=True, exist_ok=True)

    manual_path = COPYRIGHT_DIR / "软件说明书.docx"
    reg_path = COPYRIGHT_DIR / "登记信息表.docx"

    build_manual_doc().save(manual_path)
    build_registration_doc().save(reg_path)

    front_path = build_source_doc("源程序-前30页.txt", "源程序-前30页.docx")
    back_path = build_source_doc("源程序-后30页.txt", "源程序-后30页.docx")

    print("Generated Word documents:")
    print(f"  {manual_path}")
    print(f"  {reg_path}")
    print(f"  {front_path}")
    print(f"  {back_path}")


if __name__ == "__main__":
    main()
